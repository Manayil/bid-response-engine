import streamlit as st
import pandas as pd
import pdfplumber
import docx
from groq import Groq
import io
import os
import ast
from dotenv import load_dotenv
from vector_engine import search_library

# 1. System Configuration
load_dotenv()
api_key = os.getenv("GROQ_API_KEY")
client = Groq(api_key=api_key)

st.set_page_config(
    page_title="Bid and Proposal Response Engine", 
    layout="wide",
    initial_sidebar_state="expanded"
)

# Initialize Session State for Workspaces and Edits
if 'workspaces' not in st.session_state:
    st.session_state['workspaces'] = {}
if 'active_rfp' not in st.session_state:
    st.session_state['active_rfp'] = None

# 2. Data Persistence Layer
@st.cache_data
def load_datasets():
    try:
        history = pd.read_csv('data/bid_history.csv')
        library = pd.read_csv('data/capability_library.csv')
        return history, library
    except Exception as e:
        st.error(f"Data Source Error: Ensure CSV files are present in the 'data' directory. {e}")
        return None, None

history_df, library_df = load_datasets()

# 3. Utility Functions
def read_document_content(file):
    """Handles multi-format ingestion (PRD 6.1)."""
    if file.name.endswith('.pdf'):
        with pdfplumber.open(file) as pdf:
            return "".join([page.extract_text() or "" for page in pdf.pages])
    elif file.name.endswith('.docx'):
        doc = docx.Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    return None

def execute_llm_query(prompt):
    """Executes high-performance LLM inference via Groq."""
    completion = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2
    )
    return completion.choices[0].message.content

def export_to_docx(content, metadata, compliance_data):
    """Generates a professional structured proposal (PRD 6.7)."""
    doc = docx.Document()
    doc.add_heading(f"Bid Response Proposal: {metadata['name']}", 0)
    
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(f"Domain: {metadata['domain']}")
    doc.add_paragraph(f"Historical Win Probability: {metadata['score']:.1f}%")
    
    doc.add_heading("Compliance Matrix", level=1)
    for item in compliance_data:
        doc.add_paragraph(f"Requirement: {item['Requirement']}\nStatus: {item['Status']}\nEvidence: {item['Evidence']}\n")
    
    doc.add_heading("Technical Narrative", level=1)
    doc.add_paragraph(content)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# 4. Sidebar: Workspace and Ingestion Layer
with st.sidebar:
    st.header("Workspace Management")
    uploaded_file = st.file_uploader("Upload RFP Document", type=["pdf", "docx"])
    
    if uploaded_file:
        rfp_name = uploaded_file.name
        if rfp_name not in st.session_state['workspaces']:
            with st.status(f"Ingesting {rfp_name}...") as status:
                raw_text = read_document_content(uploaded_file)
                
                # Metadata and Requirement Extraction (PRD 6.2)
                domains = library_df['Domain'].unique().tolist()
                meta_prompt = f"""
                Extract the following from the RFP text and provide a Python dictionary:
                1. 'domain': Choose one from {domains}
                2. 'deadline': Extract the submission date
                3. 'budget': Extract the total budget (PKR)
                4. 'reqs': List the 5 most critical technical requirements
                Text: {raw_text[:3000]}
                """
                meta_resp = execute_llm_query(meta_prompt)
                try:
                    start, end = meta_resp.find("{"), meta_resp.rfind("}") + 1
                    meta = ast.literal_eval(meta_resp[start:end])
                except:
                    meta = {"domain": "Unspecified", "deadline": "TBD", "budget": "TBD", "reqs": []}

                # Semantic Search / RAG Mapping (PRD 6.3 & 6.5)
                comp_matrix = []
                success_count = 0
                for r in meta.get('reqs', []):
                    search_res = search_library(r)
                    matched = search_res['distances'][0][0] < 1.1
                    if matched: success_count += 1
                    comp_matrix.append({
                        "Requirement": r,
                        "Status": "PASS" if matched else "FAIL",
                        "Evidence": search_res['documents'][0][0] if matched else "No internal evidence identified."
                    })

                # Weighted Scoring (PRD 6.6)
                hist_df = history_df[history_df['Sector'] == meta['domain']]
                win_rate = (len(hist_df[hist_df['Outcome'] == 'Win']) / len(hist_df)) * 100 if not hist_df.empty else 50.0
                compliance_rate = (success_count / len(meta['reqs'])) * 100 if meta['reqs'] else 0
                total_probability = (win_rate * 0.4) + (compliance_rate * 0.6)

                st.session_state['workspaces'][rfp_name] = {
                    "text": raw_text,
                    "domain": meta['domain'],
                    "deadline": meta['deadline'],
                    "budget": meta['budget'],
                    "compliance": comp_matrix,
                    "score": total_probability,
                    "history": win_rate
                }
                status.update(label="Analysis Ready", state="complete")
        st.session_state['active_rfp'] = rfp_name

    if st.session_state['workspaces']:
        st.divider()
        st.subheader("Active Workspaces")
        for key in st.session_state['workspaces'].keys():
            if st.button(key, use_container_width=True):
                st.session_state['active_rfp'] = key

# 5. Main Application Interface
st.title("Bid and Proposal Response Engine")

if st.session_state['active_rfp']:
    ws = st.session_state['workspaces'][st.session_state['active_rfp']]
    
    st.subheader(f"Current Workspace: {st.session_state['active_rfp']}")
    
    # Overview Metrics
    m1, m2, m3 = st.columns(3)
    m1.metric("Win Probability", f"{ws['score']:.1f}%")
    m2.metric("Submission Deadline", ws['deadline'])
    m3.metric("Project Budget", ws['budget'])

    tab_stats, tab_matrix, tab_draft = st.tabs(["Strategic Analytics", "Compliance Matrix", "Technical Proposal"])

    with tab_stats:
        st.write(f"This assessment is based on a {ws['history']:.1f}% historical win rate in the {ws['domain']} sector.")
        if ws['score'] >= 75:
            st.success("Decision: GO. Capability alignment and performance history are optimal.")
        elif ws['score'] >= 50:
            st.warning("Decision: REVIEW. Substantial technical gaps identified.")
        else:
            st.error("Decision: NO-GO. Strategic mismatch or lack of historical evidence.")

    with tab_matrix:
        st.subheader("Compliance Mapping Results")
        for item in ws['compliance']:
            color = "green" if item['Status'] == "PASS" else "red"
            st.markdown(f"**{item['Requirement']}**")
            st.markdown(f"<span style='color:{color}; font-weight:bold;'>{item['Status']}</span>", unsafe_allow_html=True)
            st.caption(f"Internal Evidence: {item['Evidence']}")
            st.divider()

    with tab_draft:
        st.subheader("Automated Narrative Generation")
        
        # Human-in-the-loop editing (PRD 4.1 & 6.4)
        draft_key = f"edit_{st.session_state['active_rfp']}"
        if draft_key not in st.session_state:
            with st.spinner("Drafting initial response..."):
                evidence = ws['compliance'][0]['Evidence'] if ws['compliance'] else ""
                prompt = f"Draft a professional bid narrative. Domain: {ws['domain']}. Context: {ws['text'][:1500]}. Primary Evidence: {evidence}."
                st.session_state[draft_key] = execute_llm_query(prompt)
        
        edited_narrative = st.text_area("Edit Proposal Draft", value=st.session_state[draft_key], height=400)
        st.session_state[draft_key] = edited_narrative

        # Structured Export (PRD 6.7)
        meta_for_export = {"name": st.session_state['active_rfp'], "domain": ws['domain'], "score": ws['score']}
        word_doc = export_to_docx(edited_narrative, meta_for_export, ws['compliance'])
        
        st.download_button(
            label="Export Structured Proposal (.docx)",
            data=word_doc,
            file_name=f"Proposal_{st.session_state['active_rfp']}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
else:
    st.info("System operational. Please upload an RFP via the sidebar to initiate the analysis pipeline.")

st.divider()
st.caption("Standardized Response Engine | Analysis Pipeline v1.0 | Built for CUST Hackathon")